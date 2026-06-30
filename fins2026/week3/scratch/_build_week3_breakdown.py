"""Build the plain-language Week 3 breakdown as a Word document.

Beginner-facing guide. Uses Word built-in styles (Title, Heading 1/2,
Normal) plus a light monospace box for command blocks.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH = Path(__file__).resolve().parent / "week3_breakdown.docx"

C_CODE_BG = "F2F3F4"
C_MONO = (0x1A, 0x1A, 0x2E)


def shade(paragraph, hex_color: str) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def add_code(doc: Document, lines: list[str]) -> None:
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(12)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        shade(p, C_CODE_BG)
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(*C_MONO)


def add_body(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="Normal")


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def main() -> None:
    doc = Document()

    doc.add_paragraph("Week 3: A Complete Beginner's Breakdown", style="Title")

    # The big picture
    doc.add_heading("The big picture: what Week 3 is really about", level=1)
    add_body(
        doc,
        "Week 3 teaches you one skill above all others: how to predict the future "
        "value of a number that changes over time, and how to honestly check whether "
        "your prediction is any good. The number it uses as the main teaching example "
        "is Australia's unemployment rate, measured every month.",
    )
    add_body(
        doc,
        "That is it. Everything else in the folder, the dozens of scripts, the two "
        "apps, the figures, exists to support that single idea. The week is built so "
        "you learn the thinking first using small, simple scripts, and only at the very "
        "end do you open a polished app that does all of this automatically. The course "
        "is deliberately telling you not to open the fancy app first. You build "
        "understanding step by step, then you are allowed to see the finished product.",
    )
    add_body(
        doc,
        "There are two parallel halves to the week. The main half uses Australian "
        "economic data. The second half repeats the exact same lessons using United "
        "States data. The U.S. half is optional and only meant for after you fully "
        "understand the Australian one. So for now, ignore everything with \"us\" in its "
        "name, you can come back to it later.",
    )

    # Core idea
    doc.add_heading("The core idea behind forecasting, in everyday terms", level=1)
    add_body(
        doc,
        "Imagine you want to guess next month's unemployment rate. The simplest "
        "possible guess is: it will be the same as this month. That dead-simple guess "
        "has a name in this course, the naive benchmark (a benchmark is just a baseline "
        "to beat). It sounds too simple to be useful, but it is surprisingly hard to "
        "beat, and that is exactly the point. Any fancier model you build has to prove "
        "it does better than just saying \"same as last time.\" If it cannot beat that, "
        "the fancy model is worthless. This humility is the heart of the week.",
    )
    add_body(
        doc,
        "The week also teaches an important trick: instead of predicting the "
        "unemployment rate itself, you predict the change in it from one month to the "
        "next. The reason is that the raw rate drifts slowly and trends over long "
        "periods, which makes it misleading to model directly. The month-to-month "
        "change is steadier and better behaved, so models handle it more reliably. Once "
        "you have predicted the change, you can add it back to the last known rate to "
        "get the predicted rate. The week has a formal test for deciding whether a "
        "series should be modeled as a level or as a change, the \"unit root\" or "
        "\"stationarity\" check, which is just a statistical way of asking: is this "
        "series stable enough to model directly, or do I need to look at its changes "
        "instead?",
    )

    # Testing forecasts
    doc.add_heading("How you actually test if a forecast is good", level=1)
    add_body(
        doc,
        "This is the most important practical concept in the week. You never judge a "
        "forecast by how well it fits data it has already seen, that is like grading a "
        "student on the exact questions they studied. Instead, you split the timeline "
        "in two:",
    )
    add_bullets(
        doc,
        [
            "In-sample (the training period): all data up to and including December "
            "2019. The model is allowed to learn from this.",
            "Out-of-sample (the testing period): everything from January 2020 onward. "
            "The model has never seen this, so checking its predictions here is a fair "
            "test.",
        ],
    )
    add_body(
        doc,
        "The model makes predictions for the testing period, and you compare those "
        "predictions to what actually happened. Several scorecards measure the gap "
        "between predicted and actual; you will see these abbreviations repeatedly:",
    )
    add_bullets(
        doc,
        [
            "MAE: the average size of the mistakes (how far off you were, on average).",
            "RMSE: similar, but punishes big mistakes more harshly. This is the main "
            "score used to rank models.",
            "MASE: your error compared to the naive benchmark's error. Below 1 means "
            "you beat the naive guess; above 1 means you lost to it.",
            "Out-of-sample R-squared: how much better than the naive benchmark you did, "
            "as a percentage-style figure.",
        ],
    )

    # The ladder
    doc.add_heading("The model ladder: climbing from simple to complex", level=1)
    add_body(
        doc,
        "Week 3 introduces forecasting models as a ladder, one rung at a time. Each "
        "rung adds a little more sophistication, and at every rung you re-check whether "
        "the added complexity actually earned its keep against the naive benchmark. "
        "Here is the ladder in plain terms.",
    )
    rungs = [
        ("Rung 1, Naive", "Next month equals this month. The baseline everything must beat."),
        ("Rung 2, AR (autoregression)", "Next month depends on recent past months. The "
         "model looks at how the series moved in the last one or few months and "
         "projects that pattern forward. \"AR(1)\" just means it looks back one month."),
        ("Rung 3, ARMA", "Adds a second ingredient on top of AR: it also accounts for "
         "recent random shocks or surprises, not just the recent values. It is AR with "
         "a memory of recent errors."),
        ("Rung 4, ARX", "The \"X\" means extra outside information. Until now the model "
         "only looked at the unemployment series itself. ARX adds an outside clue, "
         "first last month's change in the cash interest rate, then last month's change "
         "in commodity prices. The question being tested: does looking at the wider "
         "economy improve the unemployment forecast?"),
        ("Rung 5, ARMAX", "Combines the ARMA backbone with those outside clues. The "
         "fullest version of the classic time-series approach."),
        ("Rung 6, OLS regression", "A different style entirely. A plain regression that "
         "hand-picks a handful of sensible economic predictors (cash rate change, "
         "commodity prices, the trade-weighted exchange-rate index, labour-market "
         "participation, job vacancies) and weighs them up together."),
        ("Rung 7, Elastic Net", "Same regression idea but with a much wider list of "
         "possible predictors, plus an automatic technique (\"shrinkage\") that politely "
         "ignores the ones that do not help. This stops the model from being fooled by "
         "too many overlapping clues."),
        ("Rung 8, Horse race", "Not a new model, a contest. It runs every model above "
         "on the same fair testing period and ranks them by RMSE so you can see, head "
         "to head, who actually won."),
        ("Rung 9, Ensemble", "Tests a classic insight: averaging several decent "
         "forecasts together is often better than trusting any single one. It checks "
         "whether a simple average of the good models beats both the naive benchmark "
         "and the single best model."),
    ]
    for title, desc in rungs:
        p = doc.add_paragraph(style="Normal")
        r = p.add_run(title + ". ")
        r.bold = True
        p.add_run(desc)

    # Sequence
    doc.add_heading("The exact sequence you are meant to follow", level=1)
    add_body(
        doc,
        "The week is organized in three layers, and you do them in order. Do not skip "
        "ahead.",
    )
    p = doc.add_paragraph(style="Normal")
    p.add_run("Layer 1, the core in-class path. ").bold = True
    p.add_run(
        "This is the required, minimum route through the week. You run a short list of "
        "scripts one after another, and each one prints results and saves figures. In "
        "order: build the data series, run the stationarity check, naive, AR, ARMA, "
        "ARX, horse race, make the story figures, then open the app. That is the spine "
        "of the whole week. The first script (make_beginner_forecasting_series.py) just "
        "creates clean data files to work from; you always run it first."
    )
    p = doc.add_paragraph(style="Normal")
    p.add_run("Layer 2, the extension path. ").bold = True
    p.add_run(
        "Optional deepening. This is where ARMAX, OLS, Elastic Net, and the ensemble "
        "live, plus the entire U.S. mirror version. Only do this once Layer 1 genuinely "
        "makes sense to you. It adds depth but no new core ideas."
    )
    p = doc.add_paragraph(style="Normal")
    p.add_run("Layer 3, the finished product. ").bold = True
    p.add_run(
        "Now, and only now, you open the polished apps. There are two: the main "
        "Australia forecast monitor (app/streamlit_app.py) and a companion U.S. macro "
        "app (us_app/streamlit_app.py). These are interactive dashboards, built with a "
        "tool called Streamlit, that do everything the scripts did, but with buttons, "
        "charts, and tabs. The Australia app has views named Overview, Australia "
        "Snapshot, Forecasts, Model Comparison, Backtests, U.S. Context, Data, and "
        "Methodology. Walking through them in that order retells the whole story you "
        "just built by hand."
    )
    add_body(
        doc,
        "A practical note on the apps: they start in \"fixture mode,\" which means they "
        "run on a frozen, saved copy of the data so they work even with no internet. "
        "There is a switch in the sidebar to flip to \"live mode,\" which pulls fresh "
        "data from public sources (the Reserve Bank of Australia and the U.S. FRED "
        "database) when you are online. Beginners should stay in fixture mode.",
    )

    # How to run
    doc.add_heading("How to actually run it (macOS)", level=1)
    add_body(
        doc,
        "Everything runs from the repository's top folder using the project's own "
        "bundled Python. You do not type \"python\" by itself; you point at the "
        "project's interpreter at ./.venv/bin/python. The core Layer 1 sequence is:",
    )
    add_code(
        doc,
        [
            "./.venv/bin/python fins2026/week3/scripts/make_beginner_forecasting_series.py",
            "./.venv/bin/python fins2026/week3/scripts/run_beginner_unit_root_check.py",
            "./.venv/bin/python fins2026/week3/scripts/run_beginner_naive_forecast.py",
            "./.venv/bin/python fins2026/week3/scripts/run_beginner_ar_forecast.py",
            "./.venv/bin/python fins2026/week3/scripts/run_beginner_arma_forecast.py",
            "./.venv/bin/python fins2026/week3/scripts/run_beginner_arx_forecast.py",
            "./.venv/bin/python fins2026/week3/scripts/run_beginner_model_horse_race.py",
            "./.venv/bin/python fins2026/week3/scripts/make_beginner_forecast_story_figures.py",
        ],
    )
    add_body(
        doc,
        "Then launch the app with: streamlit run fins2026/week3/app/streamlit_app.py. "
        "Run the scripts in this order because each builds on files the previous one "
        "produced. Everything they generate is saved under fins2026/week3/results/ "
        "(tables, figures, and data). That folder is treated as disposable, it is "
        "deliberately not saved into version control, so you can re-run the scripts any "
        "time to refresh it.",
    )

    # Folder map
    doc.add_heading("Where things live, so the folder stops looking scary", level=1)
    add_bullets(
        doc,
        [
            "scripts/: the runnable step-by-step lessons. These are what you actually "
            "launch. The run_beginner_* ones are the Australian ladder; the "
            "run_us_beginner_* ones are the optional U.S. mirror.",
            "code/: the shared engine room. The scripts are thin; the real forecasting "
            "logic lives here (e.g. beginner_forecasting.py). You usually read these "
            "only if curious.",
            "data/: the original, committed input data (the Australian macro file). "
            "Treat it as read-only source material.",
            "results/: everything the scripts generate: figures, tables, forecasts. "
            "Disposable and re-runnable.",
            "app/ and us_app/: the two finished dashboards.",
            "tests/: automatic checks that confirm nothing is broken; you do not "
            "normally touch these.",
            "The guides: README.md (the map), BEGINNER_FORECASTING.md (your "
            "step-by-step teacher), WORKSHOP.md (the run sheet), DATA_GUIDE.md (what "
            "the data means), and SUBMISSION_CHECKLIST.md (for when you deploy your app "
            "at the end).",
        ],
    )

    # One sentence
    doc.add_heading("The one sentence to remember", level=1)
    p = doc.add_paragraph(style="Normal")
    run = p.add_run(
        "Pick a target, predict its monthly change, beat the \"same as last month\" "
        "benchmark, add complexity only when it earns its place out-of-sample, and "
        "trust the polished app only after you understand the simple scripts underneath "
        "it."
    )
    run.italic = True

    doc.save(OUTPUT_PATH)
    print(f"Wrote {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
